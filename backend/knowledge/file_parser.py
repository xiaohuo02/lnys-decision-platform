# -*- coding: utf-8 -*-
"""backend/knowledge/file_parser.py — 多格式文件解析器

策略模式：根据文件类型分发到对应解析器。
支持: PDF / DOCX / XLSX / CSV / TXT / 图片(OCR)
"""
from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import Dict, List, Optional, Protocol

from loguru import logger


class ParseResult:
    """解析结果"""
    __slots__ = ("text", "pages", "tables", "metadata", "warnings")

    def __init__(
        self,
        text: str = "",
        pages: Optional[List[str]] = None,
        tables: Optional[List[str]] = None,
        metadata: Optional[Dict] = None,
        warnings: Optional[List[str]] = None,
    ):
        self.text = text
        self.pages = pages or []
        self.tables = tables or []
        self.metadata = metadata or {}
        self.warnings = warnings or []


class FileParserProtocol(Protocol):
    def parse(self, file_path: Path) -> ParseResult: ...


# ── 具体解析器 ────────────────────────────────────────────────────

class TxtParser:
    def parse(self, file_path: Path) -> ParseResult:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return ParseResult(text=text, metadata={"format": "txt"})


class CsvParser:
    def parse(self, file_path: Path) -> ParseResult:
        text_parts = []
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                headers = next(reader, None)
                if headers:
                    text_parts.append(" | ".join(headers))
                for row in reader:
                    text_parts.append(" | ".join(row))
        except Exception as e:
            return ParseResult(text="", warnings=[f"CSV parse error: {e}"])
        return ParseResult(
            text="\n".join(text_parts),
            tables=["\n".join(text_parts)],
            metadata={"format": "csv", "rows": len(text_parts)},
        )


class PdfParser:
    def parse(self, file_path: Path) -> ParseResult:
        try:
            import pdfplumber
        except ImportError:
            return ParseResult(text="", warnings=["pdfplumber not installed"])

        pages = []
        tables = []
        warnings = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(text)
                    # 表格提取
                    for tbl in (page.extract_tables() or []):
                        rows = [" | ".join(str(c or "") for c in row) for row in tbl]
                        tables.append("\n".join(rows))
                    # 检测扫描页（无文本 → 可能需要 OCR）
                    if not text.strip():
                        warnings.append(f"page_{i+1}_no_text_may_need_ocr")
        except Exception as e:
            warnings.append(f"PDF parse error: {e}")

        full_text = "\n\n".join(pages)
        return ParseResult(
            text=full_text,
            pages=pages,
            tables=tables,
            metadata={"format": "pdf", "page_count": len(pages)},
            warnings=warnings,
        )


class DocxParser:
    def parse(self, file_path: Path) -> ParseResult:
        try:
            from docx import Document
        except ImportError:
            return ParseResult(text="", warnings=["python-docx not installed"])

        try:
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 表格提取
            tables = []
            for tbl in doc.tables:
                rows = []
                for row in tbl.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables.append("\n".join(rows))

            full_text = "\n".join(paragraphs)
            return ParseResult(
                text=full_text,
                tables=tables,
                metadata={"format": "docx", "paragraphs": len(paragraphs)},
            )
        except Exception as e:
            return ParseResult(text="", warnings=[f"DOCX parse error: {e}"])


class XlsxParser:
    def parse(self, file_path: Path) -> ParseResult:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return ParseResult(text="", warnings=["openpyxl not installed"])

        try:
            wb = load_workbook(str(file_path), read_only=True, data_only=True)
            all_text = []
            tables = []
            for sheet in wb.worksheets:
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    table_text = "\n".join(rows)
                    tables.append(table_text)
                    all_text.append(f"[Sheet: {sheet.title}]\n{table_text}")
            wb.close()
            return ParseResult(
                text="\n\n".join(all_text),
                tables=tables,
                metadata={"format": "xlsx", "sheets": len(all_text)},
            )
        except Exception as e:
            return ParseResult(text="", warnings=[f"XLSX parse error: {e}"])


class ImageParser:
    """图片 → OCR 文本。需要 ocr_engine 模块支持。"""
    def parse(self, file_path: Path) -> ParseResult:
        try:
            from backend.knowledge.ocr_engine import OCREngine
            engine = OCREngine.get_instance()
            text, confidence = engine.recognize_file(file_path)
            warnings = []
            if confidence < 0.3:
                warnings.append(f"ocr_low_confidence_{confidence:.2f}")
            elif confidence < 0.6:
                warnings.append(f"ocr_medium_confidence_{confidence:.2f}")
            return ParseResult(
                text=text,
                metadata={"format": "ocr", "ocr_confidence": confidence},
                warnings=warnings,
            )
        except Exception as e:
            return ParseResult(text="", warnings=[f"OCR error: {e}"])


# ── 解析器注册表 + 工厂 ──────────────────────────────────────────

_PARSER_MAP: Dict[str, FileParserProtocol] = {
    ".txt": TxtParser(),
    ".csv": CsvParser(),
    ".pdf": PdfParser(),
    ".docx": DocxParser(),
    ".xlsx": XlsxParser(),
    ".xls": XlsxParser(),
    ".jpg": ImageParser(),
    ".jpeg": ImageParser(),
    ".png": ImageParser(),
    ".bmp": ImageParser(),
    ".tiff": ImageParser(),
    ".webp": ImageParser(),
}

SUPPORTED_EXTENSIONS = set(_PARSER_MAP.keys())


def parse_file(file_path: Path) -> ParseResult:
    """根据文件扩展名选择解析器并解析。"""
    ext = file_path.suffix.lower()
    parser = _PARSER_MAP.get(ext)
    if parser is None:
        return ParseResult(
            text="",
            warnings=[f"unsupported_format:{ext}"],
            metadata={"format": ext},
        )
    logger.debug(f"[FileParser] parsing {file_path.name} with {type(parser).__name__}")
    return parser.parse(file_path)
